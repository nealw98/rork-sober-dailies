#!/usr/bin/env python3
"""
Big Book Word Document Extractor - Using Table of Contents

Uses the table of contents to identify exact page positions, then extracts
chapters based on section breaks in the Word document.

Requirements:
    pip install python-docx

Usage:
    python extract_bigbook_from_word_v3.py
"""

import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional
try:
    from docx import Document
except ImportError:
    print("Please install python-docx: pip install python-docx")
    exit(1)

# Chapter mapping from TOC
# Format: chapter_id -> (doc_page_number, book_page_number, title)
CHAPTER_MAP = {
    'preface': (None, 'xi', 'Preface'),
    'foreword-first': (None, 'xiii', 'Foreword to First Edition'),
    'foreword-second': (None, 'xv', 'Foreword to Second Edition'),
    'doctors-opinion': (None, 'xxiii', "The Doctor's Opinion"),
    'chapter-1': (12, 1, "Bill's Story"),
    'chapter-2': (None, 17, 'There Is a Solution'),
    'chapter-3': (None, 30, 'More About Alcoholism'),
    'chapter-4': (None, 44, 'We Agnostics'),
    'chapter-5': (None, 58, 'How It Works'),
    'chapter-6': (None, 72, 'Into Action'),
    'chapter-7': (None, 89, 'Working with Others'),
    'chapter-8': (None, 104, 'To Wives'),
    'chapter-9': (None, 122, 'The Family Afterward'),
    'chapter-10': (None, 136, 'To Employers'),
    'chapter-11': (None, 151, 'A Vision for You'),
}

# Page number info for each chapter
CHAPTER_INFO = {
    'preface': {'title': 'Preface', 'book_pages': ['xi', 'xii'], 'use_roman': True},
    'foreword-first': {'title': 'Foreword to First Edition', 'book_pages': ['xiii', 'xiv'], 'use_roman': True},
    'foreword-second': {'title': 'Foreword to Second Edition', 'book_pages': ['xv', 'xvi', 'xvii', 'xviii', 'xix', 'xx', 'xxi'], 'use_roman': True},
    'doctors-opinion': {'title': "The Doctor's Opinion", 'book_pages': ['xxiii', 'xxiv', 'xxv', 'xxvi', 'xxvii', 'xxviii', 'xxix'], 'use_roman': True},
    'chapter-1': {'title': "Bill's Story", 'chapter_number': 1, 'book_pages': list(range(1, 17))},
    'chapter-2': {'title': 'There Is a Solution', 'chapter_number': 2, 'book_pages': list(range(17, 30))},
    'chapter-3': {'title': 'More About Alcoholism', 'chapter_number': 3, 'book_pages': list(range(30, 44))},
    'chapter-4': {'title': 'We Agnostics', 'chapter_number': 4, 'book_pages': list(range(44, 58))},
    'chapter-5': {'title': 'How It Works', 'chapter_number': 5, 'book_pages': list(range(58, 72))},
    'chapter-6': {'title': 'Into Action', 'chapter_number': 6, 'book_pages': list(range(72, 89))},
    'chapter-7': {'title': 'Working with Others', 'chapter_number': 7, 'book_pages': list(range(89, 104))},
    'chapter-8': {'title': 'To Wives', 'chapter_number': 8, 'book_pages': list(range(104, 122))},
    'chapter-9': {'title': 'The Family Afterward', 'chapter_number': 9, 'book_pages': list(range(122, 136))},
    'chapter-10': {'title': 'To Employers', 'chapter_number': 10, 'book_pages': list(range(136, 151))},
    'chapter-11': {'title': 'A Vision for You', 'chapter_number': 11, 'book_pages': list(range(151, 165))},
}

def find_chapter_starts_by_section_breaks(doc: Document) -> Dict[str, int]:
    """
    Find chapter start positions by looking for chapter titles in the actual content.
    Skip the table of contents (first ~50 paragraphs).
    """
    print("\n🔍 Finding chapter starts in document...")
    
    # Map chapter titles to paragraph indices where they appear
    chapter_starts = {}
    
    # Start searching after the TOC (around paragraph 50)
    start_search = 50
    
    for chapter_id, (doc_page, book_page, title) in CHAPTER_MAP.items():
        title_upper = title.upper()
        
        for i in range(start_search, len(doc.paragraphs)):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # Look for the chapter title as a standalone all-caps line
            if text.upper() == title_upper:
                chapter_starts[chapter_id] = i
                print(f"   ✅ Found {title} at paragraph {i}")
                break
    
    return chapter_starts

def extract_chapter_paragraphs(doc: Document, start_idx: int, end_idx: Optional[int]) -> List[str]:
    """Extract paragraphs for a chapter between start and end indices."""
    paragraphs = []
    
    if end_idx is None:
        end_idx = len(doc.paragraphs)
    
    for i in range(start_idx, end_idx):
        if i >= len(doc.paragraphs):
            break
        
        text = doc.paragraphs[i].text.strip()
        
        # Skip empty paragraphs
        if not text:
            continue
        
        # Skip section break markers
        if 'Section Break' in text:
            continue
        
        # Skip isolated page numbers
        if re.match(r'^[xivlcdm]+$', text, re.IGNORECASE) or re.match(r'^\d+$', text):
            continue
        
        # Skip chapter titles
        if text.upper() == text and len(text) < 60:
            continue
        
        # Skip very short lines
        if len(text) < 15:
            continue
        
        paragraphs.append(text)
    
    return paragraphs

def generate_typescript_file(chapter_id: str, paragraphs: List[str]) -> str:
    """Generate TypeScript file content for a chapter."""
    chapter_info = CHAPTER_INFO[chapter_id]
    title = chapter_info['title']
    book_pages = chapter_info['book_pages']
    chapter_number = chapter_info.get('chapter_number')
    
    # Determine page range for display
    if isinstance(book_pages[0], str):
        page_range_str = f"{book_pages[0]}–{book_pages[-1]}"
        first_page = book_pages[0]
        last_page = book_pages[-1]
    else:
        page_range_str = f"{book_pages[0]}–{book_pages[-1]}"
        first_page = book_pages[0]
        last_page = book_pages[-1]
    
    # Start building the TypeScript file
    lines = [
        "import { BigBookChapter } from '@/types/bigbook-v2';",
        "",
        "/**",
        f" * {title}",
    ]
    
    if chapter_number:
        lines.append(f" * Chapter {chapter_number}")
    
    lines.extend([
        f" * ",
        f" * Pages {page_range_str}",
        " */",
        "",
        f"export const {chapter_id.replace('-', '_')}: BigBookChapter = {{",
        f"  id: '{chapter_id}',",
        f"  title: '{title.replace(chr(39), chr(92) + chr(39))}',"
    ])
    
    if chapter_number:
        lines.append(f"  chapterNumber: {chapter_number},")
    
    lines.extend([
        f"  pageRange: [{repr(first_page)}, {repr(last_page)}],",
        "  paragraphs: ["
    ])
    
    # Distribute paragraphs across pages
    total_pages = len(book_pages)
    paras_per_page = max(1, len(paragraphs) // total_pages)
    
    for i, para in enumerate(paragraphs):
        # Estimate which page this paragraph belongs to
        page_index = min(i // paras_per_page, total_pages - 1)
        page_num = book_pages[page_index]
        
        # Escape content for TypeScript string
        content = (para
                   .replace('\\', '\\\\')
                   .replace("'", "\\'")
                   .replace('\n', ' ')
                   .replace('\r', ''))
        
        # Handle page number (could be roman numeral or number)
        if isinstance(page_num, str):
            page_num_for_ts = f"'{page_num}'"
        else:
            page_num_for_ts = str(page_num)
        
        lines.append("    {")
        lines.append(f"      id: '{chapter_id}-p{i+1}',")
        lines.append(f"      chapterId: '{chapter_id}',")
        lines.append(f"      pageNumber: {page_num_for_ts},")
        lines.append(f"      order: {i+1},")
        lines.append(f"      content: '{content}',")
        lines.append("    }," if i < len(paragraphs) - 1 else "    }")
    
    lines.extend([
        "  ],",
        "};",
        ""
    ])
    
    return '\n'.join(lines)

def main():
    """Main extraction process."""
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    docx_path = project_root / 'BigBook2.docx'
    output_dir = project_root / 'constants' / 'bigbook-v2' / 'content'
    
    if not docx_path.exists():
        print(f"❌ Word document not found: {docx_path}")
        return
    
    print(f"📖 Reading Word document: {docx_path.name}")
    doc = Document(str(docx_path))
    
    # Find where each chapter starts
    chapter_starts = find_chapter_starts_by_section_breaks(doc)
    
    # Sort chapters by their position in the document
    sorted_chapters = sorted(chapter_starts.items(), key=lambda x: x[1])
    
    print(f"\n📊 Found {len(sorted_chapters)} chapters")
    
    # Process each chapter
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print("\n" + "=" * 70)
    print("EXTRACTING CHAPTERS")
    print("=" * 70)
    
    for idx, (chapter_id, start_idx) in enumerate(sorted_chapters):
        # Determine end index (start of next chapter or end of document)
        if idx < len(sorted_chapters) - 1:
            end_idx = sorted_chapters[idx + 1][1]
        else:
            end_idx = None
        
        print(f"\n📝 {chapter_id}: {CHAPTER_INFO[chapter_id]['title']}")
        
        # Extract paragraphs
        paragraphs = extract_chapter_paragraphs(doc, start_idx + 1, end_idx)  # +1 to skip title
        
        if not paragraphs:
            print(f"   ⚠️  WARNING: No paragraphs extracted")
            continue
        
        print(f"   ✅ Extracted {len(paragraphs)} paragraphs")
        
        # Generate TypeScript file
        ts_content = generate_typescript_file(chapter_id, paragraphs)
        
        # Write file
        output_file = output_dir / f'{chapter_id}.ts'
        output_file.write_text(ts_content, encoding='utf-8')
        print(f"   💾 Written to {output_file.name}")
    
    print("\n" + "=" * 70)
    print("✅ EXTRACTION COMPLETE!")
    print("=" * 70)
    print(f"\n📁 Generated files in: {output_dir}")
    print("\n📋 Next steps:")
    print("   1. Review one or two files to verify clean text")
    print("   2. Run validation: python scripts/validate_bigbook.py")
    print("   3. Test in your app")
    print()

if __name__ == '__main__':
    main()

