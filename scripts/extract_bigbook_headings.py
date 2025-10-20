#!/usr/bin/env python3
"""
Big Book Extractor - Heading-Based Approach

Finds actual chapter headings in the document and extracts content
between them, ensuring proper chapter boundaries.
"""

from pathlib import Path
from typing import List, Dict, Tuple
try:
    from docx import Document
except ImportError:
    print("Please install python-docx: pip install python-docx")
    exit(1)

# Chapter information with heading patterns to match
CHAPTERS = [
    {
        'id': 'preface',
        'title': 'Preface',
        'heading': 'PREFACE',
        'pages': ('xi', 'xii'),
        'expected_start': 'THIS is the',  # How it should start
    },
    {
        'id': 'foreword-first',
        'title': 'Foreword to First Edition',
        'heading': 'FOREWORD TO THE FIRST EDITION',
        'pages': ('xiii', 'xiv'),
        'expected_start': 'WE, OF Alcoholics Anonymous',
    },
    {
        'id': 'foreword-second',
        'title': 'Foreword to Second Edition',
        'heading': None,  # No heading, starts with "SINCE the original"
        'pages': ('xv', 'xxii'),
        'expected_start': 'SINCE the original Foreword',
    },
    {
        'id': 'doctors-opinion',
        'title': "The Doctor's Opinion",
        'heading': "THE DOCTOR'S OPINION",
        'pages': ('xxiii', 'xxx'),
        'expected_start': "THE DOCTOR'S OPINION",
    },
    {
        'id': 'chapter-1',
        'title': "Bill's Story",
        'heading': "BILL'S STORY",
        'pages': (1, 16),
        'number': 1,
    },
    {
        'id': 'chapter-2',
        'title': 'There Is a Solution',
        'heading': 'THERE IS A SOLUTION',
        'pages': (17, 29),
        'number': 2,
    },
    {
        'id': 'chapter-3',
        'title': 'More About Alcoholism',
        'heading': 'MORE ABO UT ALCOHOLISM',  # OCR artifact
        'pages': (30, 43),
        'number': 3,
    },
    {
        'id': 'chapter-4',
        'title': 'We Agnostics',
        'heading': 'WE AGNOSTICS',
        'pages': (44, 57),
        'number': 4,
    },
    {
        'id': 'chapter-5',
        'title': 'How It Works',
        'heading': 'HOW IT WORKS',
        'pages': (58, 71),
        'number': 5,
    },
    {
        'id': 'chapter-6',
        'title': 'Into Action',
        'heading': 'INTO ACTION',
        'pages': (72, 88),
        'number': 6,
    },
    {
        'id': 'chapter-7',
        'title': 'Working with Others',
        'heading': 'WORKING WITH OTHERS',
        'pages': (89, 103),
        'number': 7,
    },
    {
        'id': 'chapter-8',
        'title': 'To Wives',
        'heading': 'TO WIVES',
        'pages': (104, 121),
        'number': 8,
    },
    {
        'id': 'chapter-9',
        'title': 'The Family Afterward',
        'heading': 'THE FAMILY AFTERWARD',
        'pages': (122, 135),
        'number': 9,
    },
    {
        'id': 'chapter-10',
        'title': 'To Employers',
        'heading': 'TO EMPLOYERS',
        'pages': (136, 150),
        'number': 10,
    },
    {
        'id': 'chapter-11',
        'title': 'A Vision for You',
        'heading': 'A VISION FOR YOU',
        'pages': (151, 164),
        'number': 11,
    },
]

# Roman numeral map
ROMAN_MAP = {
    'xi': 11, 'xii': 12, 'xiii': 13, 'xiv': 14, 'xv': 15, 'xvi': 16, 'xvii': 17,
    'xviii': 18, 'xix': 19, 'xx': 20, 'xxi': 21, 'xxii': 22, 'xxiii': 23,
    'xxiv': 24, 'xxv': 25, 'xxvi': 26, 'xxvii': 27, 'xxviii': 28, 'xxix': 29, 'xxx': 30
}

def find_chapter_boundaries(doc: Document) -> List[Tuple[str, int, int]]:
    """Find start and end paragraph indices for each chapter."""
    boundaries = []
    
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    
    for i, chapter in enumerate(CHAPTERS):
        start_idx = None
        
        # Find chapter start
        if chapter['heading']:
            # Look for the heading
            for j, para in enumerate(paragraphs):
                if para == chapter['heading']:
                    start_idx = j + 1  # Start after the heading
                    break
        else:
            # For foreword-second, look for expected start text
            for j, para in enumerate(paragraphs):
                if chapter['expected_start'] in para:
                    start_idx = j
                    break
        
        if start_idx is None:
            print(f"   ⚠️  Could not find start of {chapter['id']}")
            continue
        
        # Find chapter end (start of next chapter or end of document)
        end_idx = len(paragraphs)
        if i + 1 < len(CHAPTERS):
            next_chapter = CHAPTERS[i + 1]
            if next_chapter['heading']:
                for j in range(start_idx + 1, len(paragraphs)):
                    if paragraphs[j] == next_chapter['heading']:
                        end_idx = j
                        break
            else:
                for j in range(start_idx + 1, len(paragraphs)):
                    if next_chapter['expected_start'] in paragraphs[j]:
                        end_idx = j
                        break
        
        boundaries.append((chapter['id'], start_idx, end_idx))
        print(f"   ✅ {chapter['title']}: paragraphs {start_idx}-{end_idx}")
    
    return boundaries

def extract_chapter_content(paragraphs: List[str], start: int, end: int) -> List[str]:
    """Extract and clean chapter content."""
    content = []
    
    for para in paragraphs[start:end]:
        if not para:
            continue
        
        # Skip very short lines
        if len(para) < 15:
            continue
        
        # Skip all-caps headings
        if para.isupper() and len(para) < 60:
            continue
        
        # Skip roman numerals and page numbers
        if para.lower() in ROMAN_MAP or para.isdigit():
            continue
        
        content.append(para)
    
    return content

def generate_typescript(chapter_info: Dict, content: List[str]) -> str:
    """Generate TypeScript file."""
    chapter_id = chapter_info['id']
    title = chapter_info['title']
    start_page, end_page = chapter_info['pages']
    chapter_num = chapter_info.get('number')
    
    is_roman = isinstance(start_page, str)
    
    # Build page list
    if is_roman:
        start_idx = list(ROMAN_MAP.keys()).index(start_page)
        end_idx = list(ROMAN_MAP.keys()).index(end_page)
        page_list = list(ROMAN_MAP.keys())[start_idx:end_idx+1]
    else:
        page_list = list(range(start_page, end_page + 1))
    
    lines = [
        "import { BigBookChapter } from '@/types/bigbook-v2';",
        "",
        "/**",
        f" * {title}",
    ]
    
    if chapter_num:
        lines.append(f" * Chapter {chapter_num}")
    
    lines.extend([
        f" * ",
        f" * Pages {start_page}–{end_page}",
        " */",
        "",
        f"export const {chapter_id.replace('-', '_')}: BigBookChapter = {{",
        f"  id: '{chapter_id}',",
    ])
    
    # Escape title
    escaped_title = title.replace("'", "\\'")
    lines.append(f"  title: '{escaped_title}',")
    
    if chapter_num:
        lines.append(f"  chapterNumber: {chapter_num},")
    
    lines.extend([
        f"  pageRange: [{repr(start_page)}, {repr(end_page)}],",
        "  paragraphs: ["
    ])
    
    # Distribute content across pages
    paras_per_page = max(1, len(content) // len(page_list)) if content else 1
    
    for i, para in enumerate(content):
        page_idx = min(i // paras_per_page, len(page_list) - 1)
        page_num = page_list[page_idx]
        
        # Escape content properly
        escaped = (para
                   .replace('\\', '\\\\')
                   .replace("'", "\\'")
                   .replace('\n', ' ')
                   .replace('\r', ''))
        
        page_num_ts = f"'{page_num}'" if isinstance(page_num, str) else str(page_num)
        
        lines.append("    {")
        lines.append(f"      id: '{chapter_id}-p{i+1}',")
        lines.append(f"      chapterId: '{chapter_id}',")
        lines.append(f"      pageNumber: {page_num_ts},")
        lines.append(f"      order: {i+1},")
        lines.append(f"      content: '{escaped}',")
        lines.append("    }," if i < len(content) - 1 else "    }")
    
    lines.extend([
        "  ],",
        "};",
        ""
    ])
    
    return '\n'.join(lines)

def main():
    """Main extraction."""
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    docx_path = project_root / 'BigBook2.docx'
    output_dir = project_root / 'constants' / 'bigbook-v2' / 'content'
    
    if not docx_path.exists():
        print(f"❌ Word document not found: {docx_path}")
        return
    
    print(f"📖 Reading: {docx_path.name}")
    doc = Document(str(docx_path))
    
    print("\n🔍 Finding chapter boundaries...")
    boundaries = find_chapter_boundaries(doc)
    
    if len(boundaries) != len(CHAPTERS):
        print(f"\n⚠️  Warning: Found {len(boundaries)} of {len(CHAPTERS)} chapters")
    
    # Get all paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    
    print("\n" + "=" * 70)
    print("EXTRACTING CHAPTERS")
    print("=" * 70)
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    for chapter_id, start_idx, end_idx in boundaries:
        chapter_info = next(c for c in CHAPTERS if c['id'] == chapter_id)
        
        print(f"\n📝 {chapter_info['title']}")
        
        # Extract content
        content = extract_chapter_content(paragraphs, start_idx, end_idx)
        
        if not content:
            print(f"   ⚠️  No content extracted")
            continue
        
        print(f"   ✅ Extracted {len(content)} paragraphs")
        
        # Show first line for verification
        if content:
            preview = content[0][:70] + "..." if len(content[0]) > 70 else content[0]
            print(f"   📄 Starts: {preview}")
        
        # Generate TypeScript
        ts_content = generate_typescript(chapter_info, content)
        
        # Write file
        output_file = output_dir / f'{chapter_id}.ts'
        output_file.write_text(ts_content, encoding='utf-8')
        print(f"   💾 Written to {output_file.name}")
    
    print("\n" + "=" * 70)
    print("✅ EXTRACTION COMPLETE!")
    print("=" * 70)
    print(f"\n📁 Files in: {output_dir}")
    print()

if __name__ == '__main__':
    main()

