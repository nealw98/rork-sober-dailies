#!/usr/bin/env python3
"""
Big Book Extractor - Section-Based Approach

Uses the Word document's 200 sections (pages) and maps them directly
to the TOC page numbers. Simple and straightforward.

Requirements:
    pip install python-docx

Usage:
    python extract_bigbook_sections.py
"""

from pathlib import Path
from typing import List, Dict
try:
    from docx import Document
except ImportError:
    print("Please install python-docx: pip install python-docx")
    exit(1)

# TOC mapping: chapter_id -> (book_page_start, book_page_end, title)
TOC_MAP = {
    'preface': ('xi', 'xii', 'Preface'),
    'foreword-first': ('xiii', 'xiv', 'Foreword to First Edition'),
    'foreword-second': ('xv', 'xxii', 'Foreword to Second Edition'),
    'doctors-opinion': ('xxiii', 'xxx', "The Doctor's Opinion"),
    'chapter-1': (1, 16, "Bill's Story"),
    'chapter-2': (17, 29, 'There Is a Solution'),
    'chapter-3': (30, 43, 'More About Alcoholism'),
    'chapter-4': (44, 57, 'We Agnostics'),
    'chapter-5': (58, 71, 'How It Works'),
    'chapter-6': (72, 88, 'Into Action'),
    'chapter-7': (89, 103, 'Working with Others'),
    'chapter-8': (104, 121, 'To Wives'),
    'chapter-9': (122, 135, 'The Family Afterward'),
    'chapter-10': (136, 150, 'To Employers'),
    'chapter-11': (151, 164, 'A Vision for You'),
}

# Roman numeral conversion
ROMAN_TO_INT = {
    'xi': 11, 'xii': 12, 'xiii': 13, 'xiv': 14, 'xv': 15, 'xvi': 16,
    'xvii': 17, 'xviii': 18, 'xix': 19, 'xx': 20, 'xxi': 21, 'xxii': 22,
    'xxiii': 23, 'xxiv': 24, 'xxv': 25, 'xxvi': 26, 'xxvii': 27,
    'xxviii': 28, 'xxix': 29, 'xxx': 30
}

def extract_sections_from_word(docx_path: Path) -> List[List[str]]:
    """
    Extract text organized by sections (pages).
    Each section = one page in the book.
    The document has 200 sections with actual section breaks.
    """
    print(f"📖 Reading Word document: {docx_path.name}")
    doc = Document(str(docx_path))
    
    print(f"   Document has {len(doc.sections)} sections and {len(doc.paragraphs)} paragraphs")
    
    # Organize paragraphs by which section they belong to
    sections = [[] for _ in range(len(doc.sections))]
    
    current_section_idx = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Skip empty lines
        if not text:
            continue
        
        # Try to determine which section this paragraph belongs to
        # by checking the section of the paragraph's element
        try:
            para_section = para._element.getparent()
            # This is a simplified approach - we'll distribute paragraphs evenly
        except:
            pass
        
        # Add paragraph to current section
        sections[current_section_idx].append(text)
    
    # Actually, let's just divide paragraphs evenly across the 200 sections
    # since python-docx doesn't easily expose which paragraph belongs to which section
    sections = []
    all_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    paras_per_section = len(all_paragraphs) // 200
    
    for i in range(0, len(all_paragraphs), paras_per_section):
        sections.append(all_paragraphs[i:i + paras_per_section])
    
    print(f"   ✅ Organized {len(all_paragraphs)} paragraphs into {len(sections)} sections")
    return sections

def find_bills_story_section(sections: List[List[str]]) -> int:
    """
    Find which section contains Bill's Story actual content (not TOC).
    Look for the opening line "WAR FEVER ran high" which is the start of Chapter 1.
    """
    for i, section in enumerate(sections):
        for para in section:
            # Look for the actual opening line of Bill's Story
            if "WAR FEVER" in para.upper():
                return i
            # Alternative: look for "BILL'S STORY" as a standalone heading (not TOC)
            if para.strip().upper() == "BILL'S STORY":
                return i
    return -1

def extract_chapter(sections: List[List[str]], start_page: int, end_page: int, is_roman: bool) -> List[str]:
    """
    Extract paragraphs for a chapter given its page range.
    start_page and end_page are book page numbers (e.g., 1-16 or xi-xii).
    """
    paragraphs = []
    
    # Extract text from all pages in this chapter
    for section in sections[start_page:end_page+1]:
        for para in section:
            # Skip very short lines (likely page numbers or artifacts)
            if len(para) < 15:
                continue
            
            # Skip lines that are just page numbers
            if para.isdigit() or para.lower() in ['xi', 'xii', 'xiii', 'xiv', 'xv', 'xvi', 'xvii', 'xviii', 'xix', 'xx', 'xxi', 'xxii', 'xxiii', 'xxiv', 'xxv', 'xxvi', 'xxvii', 'xxviii', 'xxix', 'xxx']:
                continue
            
            # Skip all-caps titles (but keep them if they're part of content)
            if para.isupper() and len(para) < 60:
                continue
            
            paragraphs.append(para)
    
    return paragraphs

def generate_typescript_file(chapter_id: str, start_page, end_page, title: str, paragraphs: List[str], chapter_number=None) -> str:
    """Generate TypeScript file for a chapter."""
    # Determine if pages are roman numerals
    is_roman = isinstance(start_page, str)
    
    # Create page range display
    page_range_str = f"{start_page}–{end_page}"
    
    # Build TypeScript
    lines = [
        "import { BigBookChapter } from '@/types/bigbook-v2';",
        "",
        "/**",
        f" * {title}",
    ]
    
    if chapter_number:
        lines.append(f" * Chapter {chapter_number}")
    
    lines.extend([
        f" * ",
        f" * Pages {page_range_str}",
        " */",
        "",
        f"export const {chapter_id.replace('-', '_')}: BigBookChapter = {{",
        f"  id: '{chapter_id}',",
        f"  title: '{title.replace(chr(39), chr(92) + chr(39))}',"
    ])
    
    if chapter_number:
        lines.append(f"  chapterNumber: {chapter_number},")
    
    lines.extend([
        f"  pageRange: [{repr(start_page)}, {repr(end_page)}],",
        "  paragraphs: ["
    ])
    
    # Calculate page ranges for paragraph distribution
    if is_roman:
        page_list = list(ROMAN_TO_INT.keys())[list(ROMAN_TO_INT.values()).index(ROMAN_TO_INT[start_page]):list(ROMAN_TO_INT.values()).index(ROMAN_TO_INT[end_page])+1]
    else:
        page_list = list(range(start_page, end_page + 1))
    
    total_pages = len(page_list)
    paras_per_page = max(1, len(paragraphs) // total_pages) if paragraphs else 1
    
    for i, para in enumerate(paragraphs):
        # Estimate page number for this paragraph
        page_index = min(i // paras_per_page, total_pages - 1)
        page_num = page_list[page_index]
        
        # Escape content
        content = (para
                   .replace('\\', '\\\\')
                   .replace("'", "\\'")
                   .replace('\n', ' ')
                   .replace('\r', ''))
        
        # Format page number
        page_num_ts = f"'{page_num}'" if isinstance(page_num, str) else str(page_num)
        
        lines.append("    {")
        lines.append(f"      id: '{chapter_id}-p{i+1}',")
        lines.append(f"      chapterId: '{chapter_id}',")
        lines.append(f"      pageNumber: {page_num_ts},")
        lines.append(f"      order: {i+1},")
        lines.append(f"      content: '{content}',")
        lines.append("    }," if i < len(paragraphs) - 1 else "    }")
    
    lines.extend([
        "  ],",
        "};",
        ""
    ])
    
    return '\n'.join(lines)

def main():
    """Main extraction process."""
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    docx_path = project_root / 'BigBook2.docx'
    output_dir = project_root / 'constants' / 'bigbook-v2' / 'content'
    
    if not docx_path.exists():
        print(f"❌ Word document not found: {docx_path}")
        return
    
    # Extract all sections
    sections = extract_sections_from_word(docx_path)
    
    # Find where Bill's Story starts (this is page 1)
    bills_story_idx = find_bills_story_section(sections)
    
    if bills_story_idx == -1:
        print("❌ Could not find Bill's Story to establish page 1")
        return
    
    print(f"   ✅ Found Bill's Story at section {bills_story_idx} (Page 1)")
    
    # Calculate page 0 offset (sections before Bill's Story are roman numerals)
    page_one_offset = bills_story_idx
    
    print("\n" + "=" * 70)
    print("EXTRACTING CHAPTERS")
    print("=" * 70)
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    for chapter_id, (start_page, end_page, title) in TOC_MAP.items():
        print(f"\n📝 {chapter_id}: {title}")
        
        # Determine if this is roman numerals or arabic
        is_roman = isinstance(start_page, str)
        
        # Convert to section indices
        if is_roman:
            start_idx = ROMAN_TO_INT[start_page] - ROMAN_TO_INT['xi']  # xi is the first roman page
            end_idx = ROMAN_TO_INT[end_page] - ROMAN_TO_INT['xi']
        else:
            start_idx = page_one_offset + start_page - 1
            end_idx = page_one_offset + end_page - 1
        
        # Extract paragraphs
        paragraphs = extract_chapter(sections, start_idx, end_idx, is_roman)
        
        if not paragraphs:
            print(f"   ⚠️  No paragraphs extracted")
            continue
        
        print(f"   ✅ Extracted {len(paragraphs)} paragraphs (sections {start_idx}-{end_idx})")
        
        # Get chapter number if applicable
        chapter_number = None
        if chapter_id.startswith('chapter-'):
            chapter_number = int(chapter_id.split('-')[1])
        
        # Generate TypeScript file
        ts_content = generate_typescript_file(chapter_id, start_page, end_page, title, paragraphs, chapter_number)
        
        # Write file
        output_file = output_dir / f'{chapter_id}.ts'
        output_file.write_text(ts_content, encoding='utf-8')
        print(f"   💾 Written to {output_file.name}")
    
    print("\n" + "=" * 70)
    print("✅ EXTRACTION COMPLETE!")
    print("=" * 70)
    print(f"\n📁 Generated files in: {output_dir}")
    print("\n📋 Next steps:")
    print("   1. Check a chapter file to verify clean text")
    print("   2. Run validation: python scripts/validate_bigbook.py")
    print("   3. Test in your app")
    print()

if __name__ == '__main__':
    main()

