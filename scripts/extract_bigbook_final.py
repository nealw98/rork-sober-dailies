#!/usr/bin/env python3
"""
Big Book Extractor - Simple Section-Based Approach

Uses the Word document's actual 200 sections (pages) and maps them to book pages.
Starts extraction from PREFACE (the actual content start).

Requirements:
    pip install python-docx

Usage:
    python extract_bigbook_final.py
"""

from pathlib import Path
from typing import List, Dict
try:
    from docx import Document
except ImportError:
    print("Please install python-docx: pip install python-docx")
    exit(1)

# TOC is just a guide - maps book pages to chapter info
CHAPTER_INFO = {
    'preface': {'pages': ('xi', 'xii'), 'title': 'Preface'},
    'foreword-first': {'pages': ('xiii', 'xiv'), 'title': 'Foreword to First Edition'},
    'foreword-second': {'pages': ('xv', 'xxii'), 'title': 'Foreword to Second Edition'},
    'doctors-opinion': {'pages': ('xxiii', 'xxx'), 'title': "The Doctor's Opinion"},
    'chapter-1': {'pages': (1, 16), 'title': "Bill's Story", 'number': 1},
    'chapter-2': {'pages': (17, 29), 'title': 'There Is a Solution', 'number': 2},
    'chapter-3': {'pages': (30, 43), 'title': 'More About Alcoholism', 'number': 3},
    'chapter-4': {'pages': (44, 57), 'title': 'We Agnostics', 'number': 4},
    'chapter-5': {'pages': (58, 71), 'title': 'How It Works', 'number': 5},
    'chapter-6': {'pages': (72, 88), 'title': 'Into Action', 'number': 6},
    'chapter-7': {'pages': (89, 103), 'title': 'Working with Others', 'number': 7},
    'chapter-8': {'pages': (104, 121), 'title': 'To Wives', 'number': 8},
    'chapter-9': {'pages': (122, 135), 'title': 'The Family Afterward', 'number': 9},
    'chapter-10': {'pages': (136, 150), 'title': 'To Employers', 'number': 10},
    'chapter-11': {'pages': (151, 164), 'title': 'A Vision for You', 'number': 11},
}

# Roman numeral lookup
ROMAN_MAP = {
    'xi': 11, 'xii': 12, 'xiii': 13, 'xiv': 14, 'xv': 15, 'xvi': 16, 'xvii': 17,
    'xviii': 18, 'xix': 19, 'xx': 20, 'xxi': 21, 'xxii': 22, 'xxiii': 23,
    'xxiv': 24, 'xxv': 25, 'xxvi': 26, 'xxvii': 27, 'xxviii': 28, 'xxix': 29, 'xxx': 30
}

def find_content_start(doc: Document) -> int:
    """Find where actual content starts (PREFACE heading)."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "PREFACE":
            return i
    return -1

def extract_paragraphs_by_section(doc: Document, start_para_idx: int) -> List[List[str]]:
    """
    Extract paragraphs organized by the document's 200 sections.
    Each section = one page.
    """
    # Get all non-empty paragraphs starting from content start
    all_paras = []
    for para in doc.paragraphs[start_para_idx:]:
        text = para.text.strip()
        if text:
            all_paras.append(text)
    
    # Distribute paragraphs across 200 sections (pages)
    sections = []
    paras_per_section = len(all_paras) // 200
    
    for i in range(0, len(all_paras), paras_per_section):
        section_paras = all_paras[i:i + paras_per_section]
        sections.append(section_paras)
    
    return sections

def get_chapter_paragraphs(sections: List[List[str]], start_section: int, end_section: int) -> List[str]:
    """Extract all paragraphs from a range of sections."""
    paragraphs = []
    
    for section_idx in range(start_section, end_section + 1):
        if section_idx >= len(sections):
            break
        
        for para in sections[section_idx]:
            # Skip very short lines (page numbers, etc.)
            if len(para) < 15:
                continue
            # Skip all-caps titles that are standalone
            if para.isupper() and len(para) < 60:
                continue
            # Skip isolated roman numerals
            if para.lower() in ROMAN_MAP:
                continue
            # Skip single numbers
            if para.isdigit():
                continue
            
            paragraphs.append(para)
    
    return paragraphs

def generate_typescript(chapter_id: str, info: Dict, paragraphs: List[str]) -> str:
    """Generate TypeScript file for a chapter."""
    start_page, end_page = info['pages']
    title = info['title']
    chapter_num = info.get('number')
    
    is_roman = isinstance(start_page, str)
    
    # Build page list
    if is_roman:
        start_idx = list(ROMAN_MAP.keys()).index(start_page)
        end_idx = list(ROMAN_MAP.keys()).index(end_page)
        page_list = list(ROMAN_MAP.keys())[start_idx:end_idx+1]
    else:
        page_list = list(range(start_page, end_page + 1))
    
    # TypeScript content
    lines = [
        "import { BigBookChapter } from '@/types/bigbook-v2';",
        "",
        "/**",
        f" * {title}",
    ]
    
    if chapter_num:
        lines.append(f" * Chapter {chapter_num}")
    
    lines.extend([
        f" * ",
        f" * Pages {start_page}–{end_page}",
        " */",
        "",
        f"export const {chapter_id.replace('-', '_')}: BigBookChapter = {{",
        f"  id: '{chapter_id}',",
        f"  title: '{title.replace(chr(39), chr(92) + chr(39))}',"
    ])
    
    if chapter_num:
        lines.append(f"  chapterNumber: {chapter_num},")
    
    lines.extend([
        f"  pageRange: [{repr(start_page)}, {repr(end_page)}],",
        "  paragraphs: ["
    ])
    
    # Distribute paragraphs across pages
    paras_per_page = max(1, len(paragraphs) // len(page_list)) if paragraphs else 1
    
    for i, para in enumerate(paragraphs):
        page_idx = min(i // paras_per_page, len(page_list) - 1)
        page_num = page_list[page_idx]
        
        content = (para
                   .replace('\\', '\\\\')
                   .replace('"', '\\"')
                   .replace("'", "\\'")
                   .replace('\n', ' ')
                   .replace('\r', ''))
        
        page_num_ts = f"'{page_num}'" if isinstance(page_num, str) else str(page_num)
        
        lines.append("    {")
        lines.append(f"      id: '{chapter_id}-p{i+1}',")
        lines.append(f"      chapterId: '{chapter_id}',")
        lines.append(f"      pageNumber: {page_num_ts},")
        lines.append(f"      order: {i+1},")
        lines.append(f"      content: '{content}',")
        lines.append("    }," if i < len(paragraphs) - 1 else "    }")
    
    lines.extend([
        "  ],",
        "};",
        ""
    ])
    
    return '\n'.join(lines)

def main():
    """Main extraction."""
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    docx_path = project_root / 'BigBook2.docx'
    output_dir = project_root / 'constants' / 'bigbook-v2' / 'content'
    
    if not docx_path.exists():
        print(f"❌ Word document not found: {docx_path}")
        return
    
    print(f"📖 Reading: {docx_path.name}")
    doc = Document(str(docx_path))
    
    # Find where actual content starts (PREFACE)
    content_start = find_content_start(doc)
    if content_start == -1:
        print("❌ Could not find PREFACE")
        return
    
    print(f"   ✅ Found content start at paragraph {content_start}")
    
    # Extract paragraphs organized by 200 sections
    sections = extract_paragraphs_by_section(doc, content_start)
    print(f"   ✅ Organized into {len(sections)} sections")
    
    # PREFACE starts at section 0 (which is roman numeral page xi)
    # So section 0 = page xi, section 1 = page xii, etc.
    
    print("\n" + "=" * 70)
    print("EXTRACTING CHAPTERS")
    print("=" * 70)
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    for chapter_id, info in CHAPTER_INFO.items():
        start_page, end_page = info['pages']
        
        # Convert book pages to section indices
        if isinstance(start_page, str):
            # Roman numerals: section 0 = page xi (roman 11)
            start_section = ROMAN_MAP[start_page] - 11  # xi is position 0
            end_section = ROMAN_MAP[end_page] - 11
        else:
            # Arabic: need to find where page 1 starts
            # Roman pages go from xi (11) to xxx (30) = 20 pages = sections 0-19
            # So page 1 = section 20
            start_section = 20 + start_page - 1
            end_section = 20 + end_page - 1
        
        print(f"\n📝 {chapter_id}: {info['title']}")
        print(f"   Pages {start_page}–{end_page} → Sections {start_section}–{end_section}")
        
        # Extract paragraphs
        paragraphs = get_chapter_paragraphs(sections, start_section, end_section)
        
        if not paragraphs:
            print(f"   ⚠️  No paragraphs extracted")
            continue
        
        print(f"   ✅ Extracted {len(paragraphs)} paragraphs")
        
        # Generate TypeScript
        ts_content = generate_typescript(chapter_id, info, paragraphs)
        
        # Write file
        output_file = output_dir / f'{chapter_id}.ts'
        output_file.write_text(ts_content, encoding='utf-8')
        print(f"   💾 Written to {output_file.name}")
    
    print("\n" + "=" * 70)
    print("✅ EXTRACTION COMPLETE!")
    print("=" * 70)
    print(f"\n📁 Files in: {output_dir}")
    print("\n📋 Next: Check chapter-1.ts to verify it starts with 'WAR FEVER'")
    print()

if __name__ == '__main__':
    main()


